from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
RESULTS = ROOT / "results" / "all_results.xlsx"
PAPER_DIR = ROOT / "paper"
FIG_DIR = PAPER_DIR / "figures"
FIG_DIR.mkdir(parents=True, exist_ok=True)
OUT = PAPER_DIR / "EECE695D_Final_Paper.docx"
FORMATTED_OUT = PAPER_DIR / "EECE695D_Final_Paper.formatted.docx"
BACKUP = PAPER_DIR / "EECE695D_Final_Paper.backup.docx"


ORDER = [
    "Full FT | S | VOC",
    "Partial FT | S | VOC",
    "LoRA | S | VOC",
    "Full FT | N | VOC",
    "Full FT | S | VisDrone",
    "Partial FT | S | VisDrone",
    "LoRA-r8 | S | VisDrone",
    "LoRA-r32 | S | VisDrone",
    "Full FT | N | VisDrone",
    "LoRA | N | VisDrone",
]

TRAINABLE_PARAMS = {
    ("Full FT", "S"): 10.18,
    ("Full FT", "N"): 4.00,
    ("Partial FT", "S"): 6.10,
    ("Partial FT", "N"): 2.40,
    ("LoRA", "S"): 0.80,
    ("LoRA", "N"): 0.32,
    ("LoRA-r8", "S"): 0.80,
    ("LoRA-r32", "S"): 10.47,
}

COLORS = {"Full FT": "#2364AA", "Partial FT": "#E07A2F", "LoRA": "#4C956C", "LoRA-r8": "#4C956C", "LoRA-r32": "#7B6D9C"}
MARKERS = {"VOC": "o", "VisDrone": "s"}
DATASET_COLORS = {"VOC": "#2E6F95", "VisDrone": "#B45F3C"}


def display_name(row):
    return f"{row['Strategy']}-{row['Model']}"


def short_name(strategy, model):
    return f"{strategy.replace('Full FT', 'Full').replace('Partial FT', 'Partial')}-{model}"


def safe_metric(row, name, default=0):
    value = row.get(name, default)
    return default if pd.isna(value) else value


def load_results():
    summary = pd.read_excel(RESULTS, "Summary")
    curves = pd.read_excel(RESULTS, "Training_Curves")
    summary["Experiment"] = pd.Categorical(summary["Experiment"], ORDER, ordered=True)
    curves["Experiment"] = pd.Categorical(curves["Experiment"], ORDER, ordered=True)
    return (
        summary.sort_values("Experiment"),
        curves.sort_values(["Experiment", "epoch"]),
    )


def setup_plot_style():
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "axes.titlesize": 12,
            "axes.labelsize": 10,
            "axes.titleweight": "bold",
            "xtick.labelsize": 9,
            "ytick.labelsize": 9,
            "legend.fontsize": 8.5,
            "figure.dpi": 180,
            "savefig.dpi": 220,
            "axes.spines.top": False,
            "axes.spines.right": False,
        }
    )


def save_loss_plot(curves):
    fig, axes = plt.subplots(1, 2, figsize=(9.6, 3.8), sharex=True)
    for ax, dataset in zip(axes, ["VOC", "VisDrone"]):
        df = curves[curves["Dataset"].eq(dataset)]
        for _, group in df.groupby("Experiment", observed=False):
            if group.empty:
                continue
            label = short_name(group["Strategy"].iloc[0], group["Model"].iloc[0])
            ax.plot(
                group["epoch"],
                group["train_loss"],
                marker="o",
                linewidth=2.0,
                markersize=4,
                label=label,
                color=COLORS.get(group["Strategy"].iloc[0], "#777777"),
                alpha=0.9,
            )
            last = group.iloc[-1]
            ax.annotate(
                f"{last['train_loss']:.1f}",
                (last["epoch"], last["train_loss"]),
                textcoords="offset points",
                xytext=(5, 0),
                fontsize=7.5,
                color=COLORS.get(group["Strategy"].iloc[0], "#555555"),
            )
        ax.set_title(dataset)
        ax.set_xlabel("Epoch")
        ax.grid(True, axis="y", alpha=0.2)
    axes[0].set_ylabel("Total training loss")
    handles, labels = axes[1].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=5, frameon=False, bbox_to_anchor=(0.5, -0.03))
    fig.suptitle("Training loss dynamics over 10 epochs", y=1.0, fontsize=14, fontweight="bold")
    fig.tight_layout(rect=(0, 0.08, 1, 0.95))
    path = FIG_DIR / "figure1_training_loss.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_ap_plot(curves):
    fig, axes = plt.subplots(1, 2, figsize=(9.6, 3.8), sharex=True)
    for ax, dataset in zip(axes, ["VOC", "VisDrone"]):
        df = curves[curves["Dataset"].eq(dataset)]
        for _, group in df.groupby("Experiment", observed=False):
            if group.empty:
                continue
            label = short_name(group["Strategy"].iloc[0], group["Model"].iloc[0])
            color = COLORS.get(group["Strategy"].iloc[0], "#777777")
            ax.plot(group["epoch"], group["AP"], marker="o", linewidth=2.0, markersize=4, label=label, color=color, alpha=0.9)
            best = group.loc[group["AP"].idxmax()]
            ax.scatter(best["epoch"], best["AP"], s=70, facecolors="white", edgecolors=color, linewidths=1.8, zorder=4)
            ax.annotate(
                f"{best['AP']:.1f}",
                (best["epoch"], best["AP"]),
                textcoords="offset points",
                xytext=(5, 5),
                fontsize=7.5,
                color=color,
            )
        ax.set_title(dataset)
        ax.set_xlabel("Epoch")
        ax.set_ylabel("AP (%)")
        ax.grid(True, axis="y", alpha=0.2)
    handles, labels = axes[1].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=5, frameon=False, bbox_to_anchor=(0.5, -0.03))
    fig.suptitle("Validation AP progression with best epoch highlighted", y=1.0, fontsize=14, fontweight="bold")
    fig.tight_layout(rect=(0, 0.08, 1, 0.95))
    path = FIG_DIR / "figure2_map_progression.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_breakdown_plot(summary):
    metrics = ["Best_AP", "AP50", "AP75", "APs", "APm", "APl"]
    labels = ["AP", "AP50", "AP75", "APs", "APm", "APl"]
    fig, axes = plt.subplots(1, 2, figsize=(9.6, 4.4), constrained_layout=True)
    for ax, dataset in zip(axes, ["VOC", "VisDrone"]):
        sub = summary[summary["Dataset"].eq(dataset)]
        matrix = sub[metrics].to_numpy(dtype=float)
        image = ax.imshow(matrix, aspect="auto", cmap="YlGnBu", vmin=0, vmax=max(90, matrix.max()))
        ax.set_yticks(range(len(sub)))
        ax.set_yticklabels([display_name(row) for _, row in sub.iterrows()])
        ax.set_title(dataset)
        ax.set_xticks(range(len(metrics)))
        ax.set_xticklabels(labels)
        for y in range(matrix.shape[0]):
            for x in range(matrix.shape[1]):
                val = matrix[y, x]
                ax.text(x, y, f"{val:.1f}", ha="center", va="center", fontsize=7.5, color="black" if val < 55 else "white")
    fig.colorbar(image, ax=axes, shrink=0.78, label="Metric value (%)")
    fig.suptitle("Best-epoch AP family metrics by strategy and object scale", y=1.04, fontsize=14, fontweight="bold")
    path = FIG_DIR / "figure3_ap_breakdown.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_detection_plot(curves):
    fig, axes = plt.subplots(1, 2, figsize=(9.6, 4.0), sharey=True)
    bar_width = 0.24
    for ax, dataset in zip(axes, ["VOC", "VisDrone"]):
        df = curves[curves["Dataset"].eq(dataset)]
        rows = []
        for _, group in df.groupby("Experiment", observed=False):
            if group.empty or group["val_f1"].isna().all():
                continue
            best = group.loc[group["val_f1"].idxmax()]
            rows.append((short_name(group["Strategy"].iloc[0], group["Model"].iloc[0]), best["val_precision"], best["val_recall"], best["val_f1"]))
        labels_local = [row[0] for row in rows]
        x = range(len(rows))
        ax.bar([i - bar_width for i in x], [row[1] for row in rows], bar_width, label="Precision", color="#486A8C")
        ax.bar(list(x), [row[2] for row in rows], bar_width, label="Recall", color="#C45A44")
        ax.bar([i + bar_width for i in x], [row[3] for row in rows], bar_width, label="F1", color="#4C956C")
        for i, row in enumerate(rows):
            ax.text(i + bar_width, row[3] + 0.015, f"{row[3]:.2f}", ha="center", fontsize=7.5)
        ax.set_title(dataset)
        ax.set_xticks(list(x))
        ax.set_xticklabels(labels_local, rotation=25, ha="right")
        ax.set_ylim(0, 0.78)
        ax.grid(True, axis="y", alpha=0.2)
    axes[0].set_ylabel("Best-epoch score")
    axes[1].legend(loc="upper left", frameon=False, fontsize=8)
    fig.suptitle("Precision, recall, and F1 at each run's best F1 epoch", y=1.0, fontsize=14, fontweight="bold")
    fig.tight_layout(rect=(0, 0, 1, 0.94))
    path = FIG_DIR / "figure4_detection_metrics.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def save_efficiency_plot(summary):
    fig, ax = plt.subplots(figsize=(7.2, 4.6))
    for _, row in summary.iterrows():
        x_val = TRAINABLE_PARAMS.get((row["Strategy"], row["Model"]), safe_metric(row, "Params_M", 1))
        color = DATASET_COLORS.get(row["Dataset"], "#555555")
        ax.scatter(
            x_val,
            row["Best_AP"],
            s=150 if row["Strategy"] in ["LoRA-r8", "Full FT"] and row["Dataset"] == "VisDrone" else 105,
            color=color,
            marker=MARKERS.get(row["Dataset"], "o"),
            edgecolors="black",
            linewidths=0.7,
            alpha=0.9,
        )
        label = f"{short_name(row['Strategy'], row['Model'])}\n{row['Best_AP']:.1f} AP"
        ax.annotate(label, (x_val, row["Best_AP"]), textcoords="offset points", xytext=(6, 4), fontsize=7.4)
    ax.set_xscale("log")
    ax.set_xlabel("Approx. trainable parameters (M, log scale)")
    ax.set_ylabel("Best AP (%)")
    ax.set_title("Accuracy-efficiency trade-off by dataset")
    ax.grid(True, which="both", alpha=0.25)
    ax.scatter([], [], marker="o", color=DATASET_COLORS["VOC"], edgecolors="black", label="VOC")
    ax.scatter([], [], marker="s", color=DATASET_COLORS["VisDrone"], edgecolors="black", label="VisDrone")
    ax.legend(frameon=False, loc="upper left")
    fig.tight_layout()
    path = FIG_DIR / "figure5_accuracy_efficiency.png"
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return path


def fit_image(path, size=(300, 210)):
    image = Image.open(path).convert("RGB")
    image.thumbnail(size, Image.Resampling.LANCZOS)
    background = Image.new("RGB", size, "white")
    background.paste(image, ((size[0] - image.width) // 2, (size[1] - image.height) // 2))
    return background


def save_qualitative_grid(summary):
    metric_lookup = {
        row["Experiment"]: f"AP {row['Best_AP']:.1f} | F1 {row['Best_F1']:.2f}" if not pd.isna(row["Best_F1"]) else f"AP {row['Best_AP']:.1f}"
        for _, row in summary.iterrows()
    }
    samples = [
        ("Full-S VOC", "Full FT | S | VOC", Path("results/full_finetune/voc_S/val_samples/9_2008_000032.webp")),
        ("Partial-S VOC", "Partial FT | S | VOC", Path("results/partial_finetune/voc_S/val_samples/9_2008_000032.webp")),
        ("LoRA-S VOC", "LoRA | S | VOC", Path("results/peft/voc_S/val_samples/9_2008_000032.webp")),
        ("Full-S VisDrone", "Full FT | S | VisDrone", Path("results/full_finetune/visdrone_S/val_samples/9_0000001_08414_d_0000013.webp")),
        ("LoRA-r32-S VisDrone", "LoRA-r32 | S | VisDrone", Path("results/peft/visdrone_S_lora_r32/val_samples/9_0000001_08414_d_0000013.webp")),
        ("LoRA-N VisDrone", "LoRA | N | VisDrone", Path("results/peft/visdrone_N_lora/val_samples/9_0000001_08414_d_0000013.webp")),
    ]
    cell_w, cell_h, label_h = 330, 230, 48
    grid = Image.new("RGB", (cell_w * 3, (cell_h + label_h) * 2), "white")
    draw = ImageDraw.Draw(grid)
    try:
        font = ImageFont.truetype("arial.ttf", 15)
        small_font = ImageFont.truetype("arial.ttf", 12)
    except OSError:
        font = ImageFont.load_default()
        small_font = ImageFont.load_default()
    for index, (label, experiment, path) in enumerate(samples):
        row, col = divmod(index, 3)
        x, y = col * cell_w, row * (cell_h + label_h)
        fill = (237, 244, 248) if "VOC" in label else (248, 240, 235)
        draw.rectangle([x, y, x + cell_w - 1, y + label_h - 1], fill=fill, outline=(200, 200, 200))
        draw.text((x + 10, y + 6), label, fill=(25, 25, 25), font=font)
        draw.text((x + 10, y + 27), metric_lookup.get(experiment, ""), fill=(75, 75, 75), font=small_font)
        if path.exists():
            grid.paste(fit_image(path, (cell_w, cell_h)), (x, y + label_h))
        else:
            draw.rectangle([x, y + label_h, x + cell_w - 1, y + label_h + cell_h - 1], outline=(180, 180, 180))
            draw.text((x + 70, y + label_h + 90), "sample missing", fill=(120, 120, 120), font=font)
    path = FIG_DIR / "figure6_qualitative_grid.png"
    grid.save(path)
    return path


def set_font(run, size=9.5, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=9.5, bold=False, italic=False, after=4, before=0, left=0, right=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.06
    if left:
        p.paragraph_format.left_indent = Inches(left)
    if right:
        p.paragraph_format.right_indent = Inches(right)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=12 if level == 1 else 10.4, bold=True, color=(31, 78, 121) if level == 1 else None)
    return p


def add_caption(doc, text):
    return paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=8.8, italic=True, after=8, before=2)


def add_image(doc, path, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=60, start=60, bottom=60, end=60):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_box(doc, title, lines):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade_cell(cell, "EEF5FA")
    set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=9.4, bold=True, color=(31, 78, 121))
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.12)
        r = p.add_run(line)
        set_font(r, size=8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return tbl


def add_table(doc, rows, headers, font_size=7.6):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = tbl.cell(0, index)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "1F4E79")
        set_cell_margins(cell, top=70, start=45, bottom=70, end=45)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_font(r, size=font_size, bold=True, color=(255, 255, 255))
    for row_index, row in enumerate(rows):
        cells = tbl.add_row().cells
        for index, value in enumerate(row):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index], top=55, start=45, bottom=55, end=45)
            if row_index % 2 == 1:
                shade_cell(cells[index], "F3F6F8")
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=font_size)
    return tbl


def build_document(summary, fig_paths):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(10)

    p = paragraph(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    r = p.add_run("Stress-Testing D-FINE: A Comparative Study of Fine-Tuning\nStrategies for Domain Adaptation in Object Detection")
    set_font(r, size=17, bold=True, color=(31, 78, 121))
    paragraph(doc, "Asefa Natnale Sitotaw", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True, after=1)
    paragraph(doc, "POSTECH, Pohang, Republic of Korea     mati2025@postech.ac.kr", align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, after=8)

    add_box(
        doc,
        "Abstract",
        [
            "This paper evaluates how D-FINE transfers to new object-detection domains under realistic hardware limits. D-FINE was selected because it is a recent DETR-family detector that improves bounding-box localization through Fine-grained Distribution Refinement, making it a strong candidate for transfer learning rather than training a detector from scratch.",
            "The study compares full fine-tuning, partial fine-tuning with a frozen backbone, and LoRA-based parameter-efficient fine-tuning on PASCAL VOC and VisDrone five-class detection tasks. Ten 10-epoch training runs were analyzed using AP, AP50, AP75, scale-specific AP, F1, precision, recall, loss curves, and qualitative validation samples. Full D-FINE-S fine-tuning is strongest on VOC with 72.69 AP, while LoRA-r8-S is best on VisDrone with 19.42 AP. The results show that full adaptation works best for natural-image transfer, but dense aerial small-object detection remains recall-limited and is not solved by simply increasing adapter capacity.",
        ],
    )

    add_box(
        doc,
        "Key Findings",
        [
            "VOC: full D-FINE-S fine-tuning is the clear winner, more than 37 AP above LoRA-S.",
            "VisDrone: LoRA-r8-S slightly outperforms full S fine-tuning, but LoRA-r32-S and LoRA-N underperform.",
            "Main bottleneck: VisDrone recall and small-object AP remain low, so the domain shift is not solved by adding adapter capacity alone.",
            "Hardware constraint: all experiments were run sequentially on one RTX 4060 8 GB GPU, so training time and memory capacity limited the number and length of experiments.",
        ],
    )

    heading(doc, "1   Background and Motivation")
    heading(doc, "1.1  Why D-FINE?", 2)
    paragraph(doc, "Object detection has moved from anchor-heavy convolutional pipelines toward transformer-based set prediction. DETR introduced end-to-end detection, and later variants improved convergence, speed, and localization. D-FINE is a strong recent member of this family because it redefines bounding-box regression as Fine-grained Distribution Refinement (FDR). Instead of predicting only a small set of coordinate values, D-FINE refines a distribution over box locations, which is useful when accurate localization is important.")
    paragraph(doc, "I chose D-FINE for this project because it offers three advantages for a transfer-learning study. First, it is already COCO-pretrained, so the experiments can ask how well strong pretrained detection features transfer to new domains. Second, it has multiple model sizes, allowing a comparison between D-FINE-S and the smaller D-FINE-N. Third, its transformer decoder makes it suitable for LoRA, which allows parameter-efficient adaptation under GPU memory and training-time limits.")
    heading(doc, "1.2  Research Question", 2)
    paragraph(doc, "The main question is: when moving D-FINE from COCO to smaller target datasets, which adaptation strategy is most useful under realistic hardware constraints? Full fine-tuning should provide maximum flexibility, but it is expensive. Partial fine-tuning reduces the update space by freezing the backbone. LoRA updates only low-rank adapter weights, which should reduce trainable parameters but may not provide enough adaptation for difficult domains.")
    paragraph(doc, "The analysis therefore studies not only final AP, but also training dynamics, AP by object scale, precision-recall behavior, parameter efficiency, and qualitative detection outputs. This flow is important because a single AP number does not explain whether a method fails from localization error, missed objects, limited model capacity, or insufficient training.")

    heading(doc, "2   Experimental Design")
    heading(doc, "2.1  Datasets and Tasks", 2)
    paragraph(doc, "Two target domains are used. PASCAL VOC represents natural-image detection and is relatively close to the visual style of COCO. VisDrone is a harder aerial dataset with crowded scenes, viewpoint changes, and many small objects. Both datasets were converted to COCO-style annotation format and reduced to five target classes so the project could be trained and evaluated repeatedly on limited hardware.")
    heading(doc, "2.2  Models", 2)
    paragraph(doc, "The experiments use D-FINE-S and D-FINE-N initialized from COCO-pretrained checkpoints. D-FINE-S is the main model because it provides stronger capacity, while D-FINE-N is used to test whether a smaller model can preserve accuracy with lower compute. This scaling comparison is especially important for VisDrone, where dense small-object detection may require more representational capacity.")
    heading(doc, "2.3  Adaptation Strategies", 2)
    paragraph(doc, "I evaluated three adaptation strategies. Full fine-tuning updates all model parameters and gives the detector the most freedom to adapt. Partial fine-tuning freezes the HGNetv2 backbone and updates the detection stack, reducing adaptation cost while keeping pretrained visual features stable. LoRA inserts low-rank adapter updates into decoder attention projections while preserving most pretrained weights. The main LoRA configuration is rank 8; additional updated VisDrone runs include rank 32 for D-FINE-S and a LoRA run for D-FINE-N.")
    heading(doc, "2.4  Hardware and Training Constraints", 2)
    paragraph(doc, "All experiments were run on a single NVIDIA GeForce RTX 4060 GPU with approximately 8 GB of VRAM. This hardware is practical for a student project but limited for transformer-based object detection. Because of memory capacity, the experiments used mixed precision (AMP), 640x640 input resolution, and a controlled 10-epoch schedule. The trainings were not run as a large parallel sweep; they had to be launched and monitored sequentially, and each additional run increased the total project time substantially.")
    paragraph(doc, "This limitation influenced the methodology. Instead of training many seeds or very long schedules, the project prioritizes a structured comparison across adaptation strategies and model sizes. The final analysis uses ten completed training logs, giving 100 epoch-level records. This is enough to compare trends, but it is still constrained by GPU capacity, total training time, and the need to repeat multiple experiments.")
    add_box(
        doc,
        "Training Protocol Summary",
        [
            "Hardware: single NVIDIA GeForce RTX 4060, 8 GB VRAM.",
            "Input and optimization: 640x640 images, AdamW, AMP enabled, validation after every epoch.",
            "Experiment scale: ten completed 10-epoch runs, analyzed from saved logs and validation samples.",
            "Metrics: COCO AP, AP50, AP75, APs/APm/APl, AR100, F1, precision, recall, total loss, and qualitative outputs.",
        ],
    )

    heading(doc, "3   Methodology")
    heading(doc, "3.1  Analysis Pipeline", 2)
    paragraph(doc, "After each training run, the saved log file was parsed into a common Excel workbook. The workbook stores a summary sheet, training curves, loss components, COCO metrics, detection metrics, AP breakdowns, and a generalization-gap placeholder. The paper figures are then regenerated from this workbook so that the result table, plots, and discussion all come from the same source.")
    heading(doc, "3.2  Evaluation Metrics", 2)
    paragraph(doc, "The main metric is AP@[0.50:0.95], which is stricter than AP50 because it averages across multiple IoU thresholds. AP50 and AP75 are included to separate loose detection from stricter localization. APs, APm, and APl are included because VisDrone contains many small objects, and small-object AP is expected to reveal the main failure mode. F1, precision, and recall are included to identify whether errors come mainly from false detections or missed detections.")
    heading(doc, "3.3  Interpreting the Results", 2)
    paragraph(doc, "The analysis follows a fixed flow. First, Table 1 compares the best epoch of every run. Second, the loss and AP curves show whether performance is still improving or has peaked. Third, the AP heatmap shows whether scale is the limiting factor. Fourth, precision-recall-F1 bars explain the detection behavior. Fifth, the parameter-efficiency plot compares accuracy against trainable parameter cost. Finally, qualitative examples check whether the visual outputs match the numerical trends.")

    heading(doc, "4   Results")
    heading(doc, "4.1  Main Quantitative Results", 2)
    paragraph(doc, "Table 1 reports the completed best-epoch metrics for every run in results/all_results.xlsx. The strongest overall result is full fine-tuning of D-FINE-S on VOC, which reaches 72.69 AP and 89.88 AP50. This result is much higher than the VOC partial and LoRA variants, meaning that full adaptation is valuable when the target domain is close to natural-image COCO-style detection.")
    paragraph(doc, "The VisDrone results tell a different story. LoRA-r8-S reaches the best VisDrone AP at 19.42, slightly above full S fine-tuning at 18.76 and partial S fine-tuning at 18.35. However, the newer LoRA-r32-S run reaches only 15.86 AP, and LoRA-N reaches only 3.68 AP. This shows that LoRA can be effective, but it is not automatically improved by higher adapter rank or a smaller backbone. The result depends on model capacity, adapter placement, and the difficulty of dense aerial small-object detection.")

    main_rows = []
    for _, row in summary.iterrows():
        main_rows.append(
            [
                short_name(row["Strategy"], row["Model"]),
                row["Model"],
                row["Dataset"],
                int(row["Best_Epoch"]),
                f"{row['Best_AP']:.2f}",
                f"{row['AP50']:.2f}",
                f"{row['AP75']:.2f}",
                f"{safe_metric(row, 'Best_F1', float('nan')):.3f}" if not pd.isna(row["Best_F1"]) else "-",
                f"{TRAINABLE_PARAMS.get((row['Strategy'], row['Model']), safe_metric(row, 'Params_M', 0)):.2f}",
                f"{row['AR100']:.2f}",
            ]
        )
    add_table(doc, main_rows, ["Run", "Model", "Data", "Epoch", "AP", "AP50", "AP75", "F1", "Train M", "AR100"], font_size=7.2)
    add_caption(doc, "Table 1: Best validation results from all completed logs. F1 is omitted for the early VOC full-S run because it predates the validator patch.")

    heading(doc, "4.2  Training Dynamics", 2)
    paragraph(doc, "Figure 1 shows that most runs improve quickly during the first few epochs and then enter a slower refinement phase. On VOC, full D-FINE-S fine-tuning reaches the lowest final loss and also gives the highest AP, so the loss and AP trends agree. For the other VOC runs, the higher final loss matches their lower AP, suggesting that the model did not adapt as completely when the backbone was frozen or when only LoRA adapters were trained.")
    paragraph(doc, "On VisDrone, the losses are closer together across S-model strategies, and their AP values are also close. This indicates that the remaining difficulty is not only optimization loss. The dataset itself is harder because objects are small, dense, and viewed from aerial perspectives. Therefore, a lower loss does not guarantee a large AP gain unless the model also improves recall and small-object localization.")
    add_image(doc, fig_paths[0], 6.85)
    add_caption(doc, "Figure 1: Total training loss per epoch. End-point labels show final loss, making convergence differences visible without consulting the raw logs.")

    heading(doc, "4.3  AP Progression per Epoch", 2)
    paragraph(doc, "Figure 2 clarifies when each run is strongest. VOC full fine-tuning rises sharply and peaks at epoch 7, after which the curve does not improve. The VOC partial, LoRA, and N-model curves improve more slowly and finish far below full S fine-tuning. This supports the conclusion that VOC benefits from updating the full D-FINE-S model.")
    paragraph(doc, "VisDrone behaves differently. LoRA-r8-S continues improving to the end of the 10-epoch schedule, while LoRA-r32-S and LoRA-N peak around epoch 7 and then decline slightly. This makes the rank-32 result especially important: more adapter capacity did not produce better adaptation and may have made the training less stable. A longer schedule might help some VisDrone runs, but only if paired with better regularization, higher resolution, or augmentation aimed at small objects.")
    add_image(doc, fig_paths[1], 6.85)
    add_caption(doc, "Figure 2: AP@[0.50:0.95] progression. Hollow markers and value labels identify each run's best epoch.")

    heading(doc, "4.4  AP Breakdown by Object Scale", 2)
    paragraph(doc, "Figure 3 explains why the VisDrone AP values remain low. VOC full fine-tuning reaches very strong large-object performance, with 83.42 APl. In contrast, VisDrone APs remains low for every method. Even the best VisDrone run, LoRA-r8-S, reaches only 10.73 APs. This means the main weakness is not simply class prediction; it is small-object localization and recall.")
    paragraph(doc, "The LoRA-N result strengthens this interpretation. LoRA-N reaches only 1.41 APs on VisDrone, showing that the smaller model plus low-rank adaptation is not enough for crowded aerial scenes. Full N fine-tuning is better than LoRA-N, but still well below D-FINE-S, so model capacity matters when the target domain contains many small objects.")
    add_image(doc, fig_paths[2], 6.85)
    add_caption(doc, "Figure 3: Heatmap of best-epoch AP metrics. The scale-specific columns show why VisDrone remains difficult: APs is low across all strategies.")

    heading(doc, "4.5  Detection Metrics", 2)
    paragraph(doc, "Figure 4 compares precision, recall, and F1. On VisDrone, precision is consistently higher than recall. This is an important result because it means the detector is not mainly failing by producing too many false positives; it is failing by missing many objects. That pattern is expected in dense aerial imagery, where many objects are small and partially crowded.")
    paragraph(doc, "LoRA-r8-S gives the best VisDrone F1 at 0.533. Full and partial S fine-tuning are close, while LoRA-r32-S is lower at 0.480 and LoRA-N is much lower at 0.210. Therefore, the detection-level metrics agree with AP: LoRA-r8-S is the best VisDrone adaptation among the tested runs, but the margin is small and the remaining bottleneck is recall.")
    add_image(doc, fig_paths[3], 6.85)
    add_caption(doc, "Figure 4: Precision, recall, and F1 at each run's best F1 epoch. VisDrone precision exceeds recall, indicating missed small objects are the main error mode.")

    heading(doc, "4.6  Accuracy-Efficiency Trade-Off", 2)
    paragraph(doc, "Figure 5 summarizes the parameter-efficiency trade-off. On VOC, full D-FINE-S fine-tuning is worth its cost because it dramatically outperforms LoRA-S and partial fine-tuning. This is the clearest case where maximizing adaptation capacity improves accuracy.")
    paragraph(doc, "On VisDrone, the best result comes from LoRA-r8-S, which uses a much smaller trainable adapter budget than full fine-tuning. However, this should not be interpreted as LoRA always being better. The rank-32 LoRA-S run performs worse, and LoRA-N performs poorly. The more accurate conclusion is that VisDrone benefits from preserving useful pretrained features while adapting the decoder carefully, but PEFT is sensitive to model size and adapter configuration.")
    add_image(doc, fig_paths[4], 6.35)
    add_caption(doc, "Figure 5: Best AP versus approximate trainable parameters. The log x-axis separates parameter-efficient adapters from full-model updates.")

    heading(doc, "4.7  Qualitative Detection Outputs", 2)
    paragraph(doc, "Figure 6 provides qualitative evidence for the numerical results. The VOC examples are visually cleaner because objects are larger and easier to separate. The VisDrone examples contain many small targets in crowded scenes, which explains why recall and APs are low. The qualitative samples therefore support the quantitative conclusion: VisDrone is a dense small-object detection problem, and the main failure mode is missed objects rather than only poor confidence calibration.")
    add_image(doc, fig_paths[5], 6.85)
    add_caption(doc, "Figure 6: Qualitative validation samples with AP/F1 badges. The VisDrone panels emphasize dense small-object scenes where recall remains difficult.")

    heading(doc, "5   Discussion")
    paragraph(doc, "The results show that the best adaptation strategy depends on the target domain. For VOC, full D-FINE-S fine-tuning is clearly best. VOC is closer to COCO-style natural imagery, so updating the whole pretrained detector allows it to adapt strongly without losing the useful pretrained representation. Partial fine-tuning and LoRA preserve more of the original model, but that restriction also limits their ability to reach the full-S result.")
    paragraph(doc, "For VisDrone, the conclusion is more nuanced. LoRA-r8-S is the best tested run, but only by a small margin over full and partial S fine-tuning. This suggests that VisDrone's main difficulty is not solved by simply updating more weights. The aerial viewpoint, small object size, and crowding create a recall bottleneck. A carefully tuned LoRA adapter can help, probably because it adapts the decoder while preserving stable pretrained features, but the rank-32 and N-model results show that PEFT is not automatically robust.")
    paragraph(doc, "The model-size comparison is also important. D-FINE-N is acceptable on VOC relative to the frozen-backbone S runs, but it is much weaker on VisDrone. This implies that small aerial objects require both enough representation capacity and an adaptation method that does not overfit or underfit. In this project, LoRA-N was too constrained, while LoRA-r32-S did not improve over rank 8.")

    heading(doc, "6   Limitations")
    paragraph(doc, "The main limitation is hardware capacity. All runs were performed on a single RTX 4060 GPU with about 8 GB VRAM. This made the project feasible, but it also limited the training schedule, resolution, batch choices, and the number of repeated trials. Because each training run took substantial time and the experiments had to be run sequentially, the study uses one main seed and 10-epoch schedules rather than many seeds or long hyperparameter sweeps.")
    paragraph(doc, "A second limitation is dataset scope. Both VOC and VisDrone were configured as five-class tasks, which makes the experiments controlled but smaller than full benchmark training. The 640x640 resolution is also a limitation for VisDrone because small aerial objects can occupy very few pixels. Higher resolution would likely improve small-object AP, but it would require more GPU memory and longer training.")
    paragraph(doc, "A third limitation is logging consistency. The earliest VOC full-S run was completed before the validator patch, so its F1, precision, and recall are missing even though its COCO AP metrics are complete. Cityscapes zero-shot scripts are present, but no final numerical Cityscapes result is included in the collected workbook, so this paper does not claim a Cityscapes generalization result.")

    heading(doc, "7   Conclusion")
    paragraph(doc, "This project evaluated D-FINE for domain adaptation under realistic single-GPU constraints. The clearest result is that full D-FINE-S fine-tuning is the best choice for VOC, reaching 72.69 AP and 89.88 AP50. This supports the conclusion that when the target domain is close to COCO-style natural images and accuracy is the priority, full fine-tuning is worth the compute cost.")
    paragraph(doc, "For VisDrone, the best tested result is LoRA-r8-S with 19.42 AP and 35.07 AP50, slightly above full S fine-tuning. However, this result should be interpreted carefully: LoRA-r32-S and LoRA-N both perform worse, so the benefit comes from the right adapter configuration, not from LoRA in general. The APs and recall results show that dense small-object aerial detection remains the central challenge.")
    paragraph(doc, "Overall, the most accurate conclusion is: D-FINE transfers well to natural-image detection through full fine-tuning, while aerial small-object detection requires more than simply increasing trainable parameters. Future work should prioritize higher-resolution VisDrone training, longer schedules when GPU resources allow, more seeds, stronger small-object augmentation, and a broader hyperparameter search for LoRA rank and adapter placement.")

    heading(doc, "References")
    refs = [
        "[1] N. Carion et al. End-to-End Object Detection with Transformers. ECCV, 2020.",
        "[2] Y. Peng et al. D-FINE: Redefine Regression Task in DETRs as Fine-grained Distribution Refinement. ICLR, 2025.",
        "[3] Y. Zhao et al. DETRs Beat YOLOs on Real-time Object Detection. CVPR, 2024.",
        "[4] E. J. Hu et al. LoRA: Low-Rank Adaptation of Large Language Models. ICLR, 2022.",
        "[5] M. Everingham et al. The PASCAL Visual Object Classes Challenge. IJCV, 2010.",
        "[6] P. Zhu et al. Detection and Tracking Meet Drones Challenge. IEEE TPAMI, 2021.",
        "[7] M. Cordts et al. The Cityscapes Dataset for Semantic Urban Scene Understanding. CVPR, 2016.",
        "[8] T.-Y. Lin et al. Microsoft COCO: Common Objects in Context. ECCV, 2014.",
        "[9] I. Loshchilov and F. Hutter. Decoupled Weight Decay Regularization. ICLR, 2019.",
        "[10] Z. Han et al. Parameter-Efficient Fine-Tuning for Large Models: A Comprehensive Survey. arXiv:2403.14608, 2024.",
    ]
    for ref in refs:
        p = paragraph(doc, ref, size=8.4, after=2)
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)

    return doc


def main():
    summary, curves = load_results()
    setup_plot_style()
    fig_paths = [
        save_loss_plot(curves),
        save_ap_plot(curves),
        save_breakdown_plot(summary),
        save_detection_plot(curves),
        save_efficiency_plot(summary),
        save_qualitative_grid(summary),
    ]

    if OUT.exists() and not BACKUP.exists():
        import shutil

        shutil.copy2(OUT, BACKUP)

    doc = build_document(summary, fig_paths)
    try:
        doc.save(OUT)
        saved_to = OUT
    except PermissionError:
        doc.save(FORMATTED_OUT)
        saved_to = FORMATTED_OUT
        print(f"Could not overwrite {OUT}; it is probably open in Word.")
    print(f"Updated {saved_to}")
    for path in fig_paths:
        print(f"Figure: {path}")
    print(f"Rows in summary: {len(summary)}")


if __name__ == "__main__":
    main()
